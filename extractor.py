import os
import re
import fitz
import pdfplumber
from docx import Document
import pypandoc
from cleaning import clean_text, find_number

# getting the extension type of the file

file = "/Users/kshiraj/Downloads/resumes2/ResumeDakshMaheshwari.pdf"
ext = os.path.splitext(file)[1].lower()

open("output.txt", "w").close()

# extracting text from pdf files 

if ext == '.pdf':
    try:
        doc = fitz.open(file)
        with open("output.txt","a",encoding="utf8") as text:
            for page in doc:
                raw_data = page.get_text("text")
                if raw_data:
                    clean_data = clean_text(raw_data)
                    text.write(clean_data + "\n")

    except:
        with pdfplumber.open(file) as pdf:
            with open("output.txt","a",encoding="utf8") as text:
                
                for page in pdf.pages:
                    content = page.extract_text()   

                    if content:
                        text.write(content + "\n")
                    tables = page.extract_tables()

                    for table in tables:
                        for row in table:
                            row_text = " | ".join(str(cell) if cell else "" for cell in row)
                            text.write(row_text + "\n")
                
                     
# extracting data from a docx file

elif ext == '.docx':
    doc = Document(file)

    with open("output.txt","a", encoding="utf8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                f.write(row_text + "\n")


# extracting files from a doc file

elif ext == '.doc':
    doc_file = file.replace(".doc",".docx")
    
    pypandoc.convert_file(file,"docx",outputfile=doc_file)

    doc = Document(doc_file)

    with open("output.txt","a", encoding="utf8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                f.write(row_text + "\n")


else:
    print("Unsupported file type")


# reading the extracted text from the 

print("\n <--- Extracted text: ---> \n")

with open("output.txt","r", encoding="utf8") as f:
    data = f.readlines()
    for line in data:
        print(line)
print("\n <--- Extraction Ended ---> \n")


