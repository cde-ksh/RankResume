import os
import fitz
import pdfplumber
from docx import Document
import pypandoc
from cleaning import clean_text

# getting the extension type of the file
 
file = "/Users/kshiraj/Downloads/resumes2/ResumeAryanSingh.pdf"
def extract_text(file, save_output = True):

    ext = os.path.splitext(file)[1].lower()
    all_text = []

    # <----- extracting text from pdf files ----->

    if ext == '.pdf':
        try:
            with fitz.open(file) as doc:
                for page in doc:
                    raw_data = page.get_text("text")

                    if raw_data:
                        clean_data = clean_text(raw_data)
                        all_text.append(clean_data)

        except Exception as e:
            print("Fitz failed: ", e)
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    content = page.extract_text()   

                    if content:
                       clean_data = clean_text(content)
                       all_text.append(clean_data)

                    # <----- Table Handling ----->

                    tables = page.extract_tables()

                    for table in tables:
                        for row in table:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                all_text.append("[TABLE] " + row_text)
                    
                        
    # extracting data from a docx file

    elif ext == '.docx':
        doc = Document(file)
        for para in doc.paragraphs:
            text = para.text.strip()

            if text:
                all_text.append(clean_text(text)) 

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                all_text.append("[TABLE] " + row_text)


    # extracting files from a doc file

    elif ext == '.doc':
        doc_file = file.replace(".doc",".docx")
        
        try:
            pypandoc.convert_file(file,"docx",outputfile=doc_file)
        
        except Exception as e:
            print("Doc conversion failed: ", e)
            return ""

        doc = Document(doc_file)
        for para in doc.paragraphs:
            text = para.text.strip()

            if text:
                clean_data = clean_text(text)
                all_text.append(clean_data)


        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                all_text.append("[TABLE] " + row_text)


    else:
        print("Unsupported file type")
        return ""

    full_text = "\n".join(all_text)

    # saving the extracted data to ouput.txt file

    if save_output:
        with open("output.txt", "w", encoding="utf8") as f:
            f.write(full_text)


    return full_text



extract_text(file)